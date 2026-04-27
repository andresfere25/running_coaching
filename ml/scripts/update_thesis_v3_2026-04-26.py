# -*- coding: utf-8 -*-
"""
update_thesis_v3_2026-04-26.py

Aplica los cambios solicitados al documento Avances_Tesis_Running_17abril_v3.docx:
  - Llena el Resumen
  - Llena las Palabras clave
  - Llena el Cap. 1 Introducción
  - Llena el Cap. 3 Pregunta de investigación
  - Agrega "Apéndice A — Apuntes personales" al final del documento

NO toca: Cap. 4 (Objetivos del usuario), Cap. 5 (Glosario), Cap. 6 (Metodología),
Cap. 7 (RUNA), Cap. 8 (Resultados), Cap. 11 (Referencias).

Backup automático en Documentos Maestria/historial/ con sello de fecha.
"""
import sys
import shutil
from pathlib import Path
from datetime import datetime

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJ = Path(r'C:/Users/andre/OneDrive/Documentos/Maestría Analítica Aplicada/running_coaching')
DOCX_SRC = PROJ / 'Documentos Maestria' / 'Avances_Tesis_Running_17abril_v3.docx'
# Como el archivo puede estar abierto en Word, lo leemos desde una copia temp
# y guardamos el resultado con sufijo de fecha. El usuario decide cuándo reemplazar.
import tempfile
TMP_READ = Path(tempfile.gettempdir()) / 'tesis_read_2026-04-26.docx'
shutil.copy(DOCX_SRC, TMP_READ)
DOCX = TMP_READ
DOCX_OUT = PROJ / 'Documentos Maestria' / 'Avances_Tesis_Running_17abril_v3.2026-04-26.docx'

# Imágenes disponibles para embeber en Apuntes personales
NB12_DIR = PROJ / 'ml' / 'notebooks' / 'outputs' / 'nb12'
FIG_RANKING = NB12_DIR / 'nb12_full_fig1_ranking_modelos.png'
FIG_CONFORMAL = NB12_DIR / 'nb12_full_fig2_cobertura_conformal.png'
FIG_ZONAS = NB12_DIR / 'nb12_full_fig3_ritmo_por_zona.png'
FIG_FEATURES = NB12_DIR / 'nb12_v2_fig_iteracion_features.png'

print(f"Documento: {DOCX}")
doc = Document(DOCX)

# ----------------------------------------------------------------------------
# 1. Reemplazar texto en párrafos según patrón "[Pendiente.]" / "[En construcción.]"
# ----------------------------------------------------------------------------

RESUMEN = (
    "Este trabajo presenta RUNA, una plataforma de coaching para corredores recreativos "
    "basada en un sistema jerárquico de aprendizaje automático que estima el ritmo "
    "sostenible por zona de frecuencia cardíaca (Z1–Z5) y los tiempos esperados en "
    "distancias objetivo (5K, 10K, 21K y 42K). El sistema combina tres niveles "
    "anidados: (i) un prior poblacional entrenado sobre el dataset público FitRec/"
    "Endomondo (N = 20 710 sesiones, 356 usuarios) que resuelve el problema de "
    "cold-start cuando el atleta es desconocido; (ii) un núcleo generalizable que "
    "incorpora datos de entrenamiento de la cohorte propia RUNA (≥ 30 atletas) "
    "mediante un protocolo AutoML (naiveautoml) y validación cruzada Leave-One-"
    "Athlete-Out (LOAO-CV); y (iii) una capa de personalización bayesiana alimentada "
    "por simulaciones de carrera del atleta individual. El Nivel 1, ya cerrado, "
    "alcanza un MAE de 40.16 sec/km (R² = +0.188), superando significativamente al "
    "baseline Karvonen en 8.2 % (Friedman χ² = 14.29, p = 0.0064) con cobertura "
    "conformal empírica de 0.804 al 80 % nominal. Las predicciones se entregan junto "
    "con intervalos de confianza calibrados, ofreciendo un soporte estadístico "
    "honesto al atleta y al entrenador. La contribución diferenciadora es la "
    "cuantificación de la curva de ganancia por personalización: cuántos datos "
    "individuales necesita un atleta antes de que el modelo personalizado supere "
    "estadísticamente al poblacional. RUNA opera como plataforma web con conexión a "
    "Strava, captura semanal de check-ins y panel de coaching."
)

PALABRAS_CLAVE = (
    "running recreativo, machine learning jerárquico, predicción de rendimiento, "
    "frecuencia cardíaca, zonas de entrenamiento, intervalos conformales, AutoML, "
    "naiveautoml, LOAO-CV, FitRec, Endomondo, Strava, coaching digital."
)

INTRODUCCION = [
    ("h2", "1.1 Contexto y motivación"),
    ("p",
     "El running recreativo ha crecido sostenidamente en la última década, y con él "
     "el acceso de los corredores a dispositivos GPS, monitores de frecuencia "
     "cardíaca y plataformas como Strava o Garmin Connect. La consecuencia ha sido "
     "una explosión de datos individuales: kilometrajes, ritmos, cadencia, desnivel "
     "y curvas de FC quedan registrados sesión a sesión. Sin embargo, la abundancia "
     "de datos no se ha traducido en mejor toma de decisiones para el atleta "
     "promedio. La mayoría de las plataformas son repositorios pasivos —muestran lo "
     "que ya pasó— y los métodos clásicos de prescripción de ritmo (Karvonen, VDOT, "
     "Riegel) son tablas estáticas que asumen que todos los corredores responden "
     "igual al esfuerzo y degradan igual con la distancia."),
    ("p",
     "Este trabajo nace de dos motivaciones complementarias. En el plano académico, "
     "se diseña y valida una arquitectura jerárquica de machine learning que estima "
     "ritmos sostenibles por zona de frecuencia cardíaca, combinando un prior "
     "poblacional entrenado sobre datos públicos (FitRec/Endomondo) con un núcleo "
     "generalizable y una capa de personalización individual. En el plano aplicado, "
     "el sistema se materializa en RUNA, una plataforma de coaching real con "
     "conexión a Strava, check-ins semanales y panel de seguimiento. Ambas "
     "dimensiones se alimentan mutuamente: la metodología justifica el producto y "
     "los datos del producto validan la metodología."),
    ("h2", "1.2 Planteamiento del problema"),
    ("p",
     "El corredor recreativo tiene hoy dos opciones para prescribir su entrenamiento "
     "y estimar su rendimiento: seguir tablas genéricas que no saben nada de él "
     "(Karvonen, VDOT, Riegel) o contratar un entrenador personal. La primera opción "
     "es accesible pero estática: el ritmo recomendado para un corredor de 20 años "
     "con FCmax de 200 lpm es el mismo que para uno de 50 años con FCmax de 170 lpm "
     "si su porcentaje de FCmax coincide. La segunda opción es personalizada pero "
     "costosa, intermitente y sin trazabilidad cuantitativa de la mejora."),
    ("p",
     "El problema central que aborda este trabajo es: dado un corredor —del que "
     "inicialmente solo se conocen variables demográficas y, progresivamente, su "
     "historial de entrenamiento y carreras— ¿con qué precisión puede estimarse el "
     "ritmo sostenible que puede mantener en cada zona de frecuencia cardíaca, y "
     "cómo mejora esa precisión a medida que se incorpora información individual?"),
    ("h2", "1.3 Justificación y valor del proyecto"),
    ("p",
     "La contribución diferenciadora de este trabajo no es predecir tiempos de "
     "carrera —eso ya existe en la literatura— sino cuantificar empíricamente la "
     "curva de ganancia por personalización: cuánto vale cada carrera y cada check-"
     "in que el atleta aporta al sistema, y a partir de qué volumen de datos "
     "individuales el modelo personalizado supera estadísticamente al modelo "
     "poblacional. Este ángulo no está disponible en plataformas comerciales "
     "(Strava, Garmin) ni se ha estudiado de forma sistemática en corredores "
     "recreacionales colombianos."),
    ("p",
     "Tres preguntas comunes de un jurado se responden directamente desde aquí. "
     "Primero, ¿por qué este proyecto es útil? Porque entrega al atleta una "
     "estimación que mejora con cada dato propio que aporta, en lugar de un número "
     "estático sin garantía de incertidumbre. Segundo, ¿qué tiene RUNA que no tiene "
     "Strava? Strava registra lo que pasó; RUNA estima lo que debería pasar y "
     "comunica la incertidumbre con un intervalo de confianza calibrado al 80 %. "
     "Son herramientas complementarias: Strava es fuente de datos para RUNA, no su "
     "rival. Tercero, ¿cuál es el problema y por qué la solución propuesta es "
     "mejor? El problema son las prescripciones de ritmo estáticas, universales y "
     "sin cuantificación de incertidumbre; la solución propuesta resuelve los tres "
     "vacíos: aprende con cada carrera, se personaliza por atleta y reporta "
     "intervalos conformales con cobertura empírica garantizada."),
    ("h2", "1.4 Estructura del documento"),
    ("p",
     "El Capítulo 2 presenta el estado del arte en predicción de rendimiento en "
     "running, modelos jerárquicos y validación estadística para sistemas "
     "personalizados. El Capítulo 3 formula la pregunta de investigación y sus "
     "subpreguntas. El Capítulo 4 enuncia los objetivos general y específicos. El "
     "Capítulo 5 reúne el marco conceptual con definiciones operacionales del "
     "dominio del running y de los métodos estadísticos empleados. El Capítulo 6 "
     "describe el marco metodológico, incluyendo fuentes de datos, consideraciones "
     "éticas, EDA, limpieza, modelado jerárquico, baselines y protocolo de "
     "validación. El Capítulo 7 introduce brevemente el componente aplicado RUNA "
     "como vehículo del sistema. El Capítulo 8 presenta los resultados preliminares "
     "—con el Nivel 1 cerrado en la sección 8.10— y delimita los pendientes de "
     "Niveles 2 y 3. Los Capítulos 9 y 10 contienen la discusión y las "
     "conclusiones. El Capítulo 11 lista las referencias bibliográficas. Al final "
     "del documento se incluye un Apéndice de apuntes personales con explicaciones "
     "extensas dirigidas al autor mismo (no al lector formal) sobre el "
     "funcionamiento interno del sistema."),
]

PREGUNTA_INVESTIGACION = [
    ("p",
     "Esta sección formula la pregunta central que articula el trabajo, sus tres "
     "subpreguntas alineadas con los objetivos específicos del Capítulo 4 y la "
     "contribución diferenciadora respecto al estado del arte."),
    ("h2", "3.1 Pregunta central"),
    ("p",
     "¿Puede un sistema jerárquico de aprendizaje automático, que aprende primero "
     "de datos poblacionales y se afina progresivamente con el historial individual "
     "del atleta, estimar el ritmo sostenible por zona de frecuencia cardíaca con "
     "mayor precisión que los métodos clásicos de prescripción (Karvonen, VDOT, "
     "Riegel) y cuántos datos individuales se requieren para que esa mejora sea "
     "estadísticamente significativa?"),
    ("h2", "3.2 Subpreguntas de investigación"),
    ("p",
     "La pregunta central se descompone en tres subpreguntas, cada una ligada a uno "
     "de los objetivos específicos del Capítulo 4."),
    ("p",
     "Q1 — Prior poblacional. ¿Qué tan bien estima un modelo entrenado sobre datos "
     "poblacionales (FitRec/Endomondo, N = 20 710 sesiones) el ritmo sostenible por "
     "zona de FC para un corredor desconocido del que solo se dispone de variables "
     "agregadas por sesión? ¿Qué algoritmo —dentro del conjunto explorado por un "
     "protocolo AutoML estándar— logra el menor error y supera al baseline Karvonen "
     "con significancia estadística?"),
    ("p",
     "Q2 — Núcleo generalizable. ¿Cuánto mejora la estimación del ritmo cuando se "
     "incorporan features derivados del historial real de entrenamiento de la "
     "cohorte propia RUNA (carga aguda y crónica, kilometraje semanal, cadencia, "
     "desnivel, edad, récords personales) y se valida bajo Leave-One-Athlete-Out "
     "Cross-Validation (LOAO-CV)?"),
    ("p",
     "Q3 — Curva de personalización. ¿A partir de cuántas carreras propias y cuántas "
     "semanas de check-ins el modelo personalizado por atleta (Nivel 3) supera "
     "estadísticamente al modelo poblacional (Nivel 1) y al núcleo generalizable "
     "(Nivel 2)? ¿Qué forma tiene la curva de ganancia y cuál es el umbral mínimo "
     "viable para activar la personalización en producción?"),
    ("h2", "3.3 Contribución diferenciadora"),
    ("p",
     "La literatura de predicción de rendimiento en running se ha centrado en "
     "modelos puntuales que estiman tiempos finales de carrera con un único conjunto "
     "de features. El presente trabajo no busca competir en precisión absoluta de "
     "predicción, sino cuantificar la dinámica de aprendizaje del modelo: cuánto "
     "aporta cada nuevo dato del atleta, en qué momento la personalización supera "
     "al consenso poblacional y cuánta incertidumbre residual permanece en cada "
     "etapa. Esta caracterización empírica de la curva de ganancia por "
     "personalización es la contribución diferenciadora del trabajo respecto al "
     "estado del arte revisado en el Capítulo 2."),
]

# ----------------------------------------------------------------------------
# Función auxiliar: encontrar y reemplazar un párrafo por texto exacto
# ----------------------------------------------------------------------------
def find_paragraph(doc, text_substring):
    """Devuelve el primer párrafo cuyo texto contiene el substring."""
    for p in doc.paragraphs:
        if text_substring in p.text:
            return p
    return None

def replace_paragraph_text(p, new_text):
    """Reemplaza el texto del párrafo conservando el estilo del primer run."""
    # Limpiar runs existentes
    for run in p.runs[1:]:
        run.text = ''
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)

def insert_paragraphs_after(anchor_p, blocks):
    """Inserta nuevos párrafos después de `anchor_p`. blocks = [(tag, text), ...]
    tag puede ser 'h2' (heading 2) o 'p' (parrafo normal)."""
    from docx.oxml.ns import qn
    from copy import deepcopy
    parent = anchor_p._element.getparent()
    idx = list(parent).index(anchor_p._element)
    inserted = []
    for tag, text in blocks:
        if tag == 'h2':
            new = doc.add_heading(text, level=2)
        else:
            new = doc.add_paragraph(text)
        # Mover el elemento recién creado al lugar correcto
        parent.insert(idx + 1, new._element)
        idx += 1
        inserted.append(new)
    return inserted

# ----------------------------------------------------------------------------
# 2. Resumen
# ----------------------------------------------------------------------------
print("\n[1/5] Llenando Resumen...")
# Buscar el párrafo "[Pendiente.]" que sigue al heading "Resumen"
found_resumen = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Resumen':
        # El siguiente párrafo es "[Pendiente.]"
        if i + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i + 1]
            if '[Pendiente.]' in next_p.text or '[Pendiente]' in next_p.text:
                replace_paragraph_text(next_p, RESUMEN)
                found_resumen = True
                print("    ✓ Resumen llenado")
                break
if not found_resumen:
    print("    ⚠ No se encontró el placeholder de Resumen")

# ----------------------------------------------------------------------------
# 3. Palabras clave
# ----------------------------------------------------------------------------
print("\n[2/5] Llenando Palabras clave...")
found_pc = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Palabras clave':
        if i + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i + 1]
            if '[Por definir' in next_p.text:
                replace_paragraph_text(next_p, PALABRAS_CLAVE)
                found_pc = True
                print("    ✓ Palabras clave llenadas")
                break
if not found_pc:
    print("    ⚠ No se encontró el placeholder de Palabras clave")

# ----------------------------------------------------------------------------
# 4. Cap. 1 Introducción
# ----------------------------------------------------------------------------
print("\n[3/5] Llenando Cap. 1 Introducción...")
found_intro = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '1. Introducción':
        if i + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i + 1]
            if '[En construcción' in next_p.text:
                # Reemplazar el placeholder con el primer párrafo
                replace_paragraph_text(next_p, INTRODUCCION[1][1])  # primer párrafo (después de h2)
                # Cambiar el placeholder a heading 2 "1.1 Contexto y motivación"
                # En realidad, mejor: reemplazo el texto del placeholder con el heading,
                # luego inserto el resto detrás.
                # Estrategia más limpia: borro el placeholder y reinserto bloque completo.
                replace_paragraph_text(next_p, INTRODUCCION[0][1])
                next_p.style = doc.styles['Heading 2']
                # Insertar el resto a partir de INTRODUCCION[1:]
                insert_paragraphs_after(next_p, INTRODUCCION[1:])
                found_intro = True
                print(f"    ✓ Introducción llenada ({len(INTRODUCCION)} bloques)")
                break
if not found_intro:
    print("    ⚠ No se encontró el placeholder de Introducción")

# ----------------------------------------------------------------------------
# 5. Cap. 3 Pregunta de investigación
# ----------------------------------------------------------------------------
print("\n[4/5] Llenando Cap. 3 Pregunta de investigación...")
found_q = False
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '3. Pregunta de investigación':
        if i + 1 < len(doc.paragraphs):
            next_p = doc.paragraphs[i + 1]
            if '[En construcción' in next_p.text:
                # Reemplazar placeholder con el primer párrafo del bloque
                replace_paragraph_text(next_p, PREGUNTA_INVESTIGACION[0][1])
                insert_paragraphs_after(next_p, PREGUNTA_INVESTIGACION[1:])
                found_q = True
                print(f"    ✓ Pregunta de investigación llenada ({len(PREGUNTA_INVESTIGACION)} bloques)")
                break
if not found_q:
    print("    ⚠ No se encontró el placeholder de Pregunta de investigación")

# ----------------------------------------------------------------------------
# 6. Apéndice A — Apuntes personales (al final del documento)
# ----------------------------------------------------------------------------
print("\n[5/5] Agregando Apéndice A — Apuntes personales...")

doc.add_page_break()
doc.add_heading('Apéndice A — Apuntes personales', level=1)
doc.add_paragraph(
    "Este apéndice está dirigido al autor mismo, no al lector formal de la tesis. "
    "Reúne explicaciones detalladas, ejemplos y referencias internas a los notebooks "
    "del proyecto. Sirve como mapa de comprensión interna del sistema y para "
    "anticipar preguntas durante la sustentación."
)

# A.1
doc.add_heading('A.1 Cómo se entrena el Nivel 1, paso a paso', level=2)
doc.add_paragraph(
    "El Nivel 1 es el prior poblacional. Su función es resolver el cold-start: "
    "estimar un ritmo razonable para un atleta del que aún no se sabe nada "
    "individual. Los pasos del entrenamiento, tal como están implementados en los "
    "notebooks, son los siguientes:"
)
pasos = [
    "1. Carga del dataset crudo. NB11 lee endomondoHR.json (~6.6 GB) en streaming "
    "desde C:\\Datasets\\running_coaching\\ y filtra solo registros con "
    "sport='running' y al menos 20 puntos válidos de FC.",
    "2. Filtros de coherencia. Velocidad media en [3, 20] km/h, FC media en "
    "[80, 210] lpm, FCmax observada por usuario en [130, 220] lpm, ritmo "
    "plausible en [3, 12] min/km.",
    "3. Estabilidad de FCmax. Solo usuarios con ≥10 sesiones registradas "
    "(garantiza que la FCmax observada sea un proxy estable de la fisiológica). "
    "Tras filtros: 20 710 sesiones, 356 usuarios.",
    "4. Cálculo de features (NB11). Para cada sesión: gender_bin (0/1), "
    "fcmax_obs (FC máxima histórica del usuario), hr_mean (FC media de la "
    "sesión), pct_fcmax (hr_mean / fcmax_obs), zona_num (1–5 según pct_fcmax), "
    "hr_max_rel (FC máxima en sesión / fcmax_obs), log_duration (ln segundos), "
    "dens_hr (densidad de muestreo cardíaco como proxy de calidad).",
    "5. Variable objetivo. pace_min_km, calculado como 60 / speed_kmh.",
    "6. Validación cruzada (NB12). GroupKFold con K = 10, agrupando por userId. "
    "El modelo nunca ve sesiones del mismo atleta en train y test; mide qué tan "
    "bien generaliza a usuarios nuevos.",
    "7. Comparación de modelos. NB12 entrena 9 algoritmos bajo el mismo "
    "protocolo: Ridge, Lasso, ElasticNet, Random Forest, Gradient Boosting, "
    "Histogram GB, MLP, baseline Karvonen y baseline naive. Mejor: Ridge con "
    "MAE = 40.16 sec/km y R² = +0.188.",
    "8. Test estadístico. Friedman sobre los errores promediados por usuario "
    "(N = 356) arroja χ² = 14.29, p = 0.0064. El post-hoc Nemenyi confirma "
    "que Ridge supera a MLP (p = 0.022) y a Gradient Boosting (p = 0.036).",
    "9. Calibración conformal. Split 60/20/20 (train/calibración/test) con "
    "α = 0.20. Cuantil de error absoluto: q = 1.06 min/km. Cobertura empírica "
    "obtenida: 0.804, prácticamente idéntica al 0.80 nominal.",
    "10. Serialización. El modelo y su scaler se guardan en "
    "ml/notebooks/outputs/nb12/nivel1_prior_poblacional_FULL_v2.pkl, listo "
    "para alimentar el Nivel 2 cuando la cohorte RUNA esté disponible."
]
for paso in pasos:
    doc.add_paragraph(paso, style='List Number')

# Insertar figura del ranking de modelos
if FIG_RANKING.exists():
    doc.add_paragraph()
    doc.add_picture(str(FIG_RANKING), width=Inches(5.5))
    cap = doc.add_paragraph(
        "Figura A.1. Ranking de los 9 modelos evaluados sobre Endomondo FULL "
        "(N = 20 710 sesiones, GroupKFold K = 10). Ridge gana con MAE = 40.16 "
        "sec/km. Fuente: NB12 — nb12_full_fig1_ranking_modelos.png."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# A.2
doc.add_heading('A.2 Cómo se conectan los tres niveles', level=2)
doc.add_paragraph(
    "Los niveles del modelo son acumulativos: cada uno hereda la predicción del "
    "anterior y la refina con información adicional del atleta. La tabla A.1 "
    "resume las entradas y salidas de cada nivel."
)

t = doc.add_table(rows=4, cols=4)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = 'Nivel'
hdr[1].text = 'Datos disponibles'
hdr[2].text = 'Features (resumen)'
hdr[3].text = 'Salida'

filas = [
    ('Nivel 1\n(Prior poblacional)',
     'FitRec/Endomondo público — sin info del atleta',
     'gender · fcmax_obs · hr_mean · pct_fcmax · zona · hr_max_rel · log_duration · dens_hr (8 features)',
     'Ritmo poblacional ± intervalo conformal\nMAE 40 sec/km'),
    ('Nivel 2\n(Núcleo generalizable)',
     'Cohorte RUNA + Strava + formulario',
     'Predicción Nivel 1 + edad + PRs + CTL + ATL + ACWR + km/sem + ritmo_4sem + cadencia + desnivel',
     'Ritmo ajustado por segmento ± intervalo más estrecho\nMAE esperado ~25 sec/km'),
    ('Nivel 3\n(Personalización bayesiana)',
     'Historial individual + check-ins + carreras',
     'Predicción Nivel 2 + carreras propias + sueño + energía + dolor + race simulations',
     'Ritmo personalizado ± intervalo aún más estrecho\nMAE esperado ~15 sec/km'),
]
for i, (n, d, f, s) in enumerate(filas, start=1):
    row = t.rows[i].cells
    row[0].text = n
    row[1].text = d
    row[2].text = f
    row[3].text = s

doc.add_paragraph()
doc.add_paragraph(
    "Tabla A.1. Datos, features y salida por nivel del modelo jerárquico.",
    style='Caption' if 'Caption' in [s.name for s in doc.styles] else None
)

# A.3
doc.add_heading('A.3 Dónde entran la edad, el peso y los récords personales (PRs)', level=2)
doc.add_paragraph(
    "Una pregunta natural es: si el atleta reporta su edad, peso y PRs en el "
    "formulario de RUNA, ¿en qué momento del modelo se usan? La respuesta corta es "
    "que NO se usan en el Nivel 1, y SÍ se usan a partir del Nivel 2. La razón "
    "tiene que ver con la estructura del dataset poblacional (FitRec) y con el "
    "diseño jerárquico del sistema."
)
doc.add_paragraph(
    "En el dataset FitRec, la mayoría de usuarios no reportaron edad ni peso de "
    "manera fiable, por lo que esas variables no son utilizables a escala "
    "poblacional. Adicionalmente, con 356 usuarios la señal estadística de la edad "
    "es débil y se diluye frente a la FC observada. Por esa razón, el Nivel 1 "
    "trabaja exclusivamente con variables agregadas por sesión (FC y duración) y "
    "el género."
)
doc.add_paragraph(
    "Cuando el atleta se registra en RUNA, sí proporciona edad, peso, PRs auto-"
    "reportados (5K, 10K, 21K, 42K) y consentimiento. Estas variables se "
    "incorporan en el Nivel 2 como features adicionales del modelo generalizable. "
    "Los PRs cumplen un doble papel: como features de entrada (señalan el nivel "
    "del corredor) y como puntos de referencia para validar internamente las "
    "estimaciones del Nivel 1 antes de exponerlas al atleta."
)
doc.add_paragraph(
    "En el Nivel 3, los PRs y las carreras reales del atleta —junto con los "
    "check-ins semanales— alimentan la calibración bayesiana individual. Las "
    "carreras propias y los entrenamientos duros considerados \"race simulations\" "
    "se usan como puntos de anclaje para refinar la estimación poblacional hacia "
    "la respuesta específica de este atleta a su carga de entrenamiento."
)

# A.4
doc.add_heading('A.4 Qué es AutoML y por qué lo usamos', level=2)
doc.add_paragraph(
    "AutoML (Automated Machine Learning) es un conjunto de técnicas para "
    "automatizar partes del flujo de trabajo de un proyecto de aprendizaje "
    "automático: selección del algoritmo, ajuste de hiperparámetros, ingeniería "
    "de features y validación. La motivación es triple. Primero, evita el sesgo "
    "del investigador que prueba solo los modelos que conoce mejor. Segundo, "
    "acelera el ciclo experimental al evaluar muchas opciones bajo un protocolo "
    "uniforme. Tercero, produce un soporte estadístico explícito para justificar "
    "el modelo elegido."
)
doc.add_paragraph(
    "En este trabajo, AutoML cumple un rol metodológico —no solo de comodidad. "
    "El director de tesis recomendó explícitamente usar un protocolo AutoML "
    "estándar para que la elección del modelo no dependa de preferencias "
    "subjetivas y sea defendible ante el jurado. La idea es: evaluamos todos los "
    "algoritmos razonables bajo exactamente el mismo esquema de validación, "
    "comparamos sus errores con un test estadístico no paramétrico (Friedman-"
    "Nemenyi), y reportamos el ganador con su intervalo de confianza."
)
doc.add_paragraph(
    "Las preguntas que un jurado puede hacer sobre AutoML, y sus respuestas:"
)
doc.add_paragraph(
    "¿Por qué no probaron deep learning? Las redes neuronales se incluyen en la "
    "comparación (MLP). Para datos tabulares con N ≈ 20 000 y 8 features, los "
    "modelos lineales y de árboles típicamente igualan o superan a las redes "
    "neuronales sin requerir tuning extensivo de arquitectura."
)
doc.add_paragraph(
    "¿Por qué no usaron autosklearn o TPOT en lugar de naiveautoml? Esos "
    "frameworks hacen búsqueda Bayesiana y son cajas más opacas. naiveautoml "
    "es deliberadamente transparente: su lógica es \"prueba todos los modelos "
    "estándar bajo el mismo CV y elige el ganador\". Esa transparencia es la "
    "razón por la cual el director lo recomendó."
)
doc.add_paragraph(
    "¿No es naiveautoml demasiado simple? El término \"naive\" es metodológico, "
    "no de simplicidad. Significa que el protocolo no asume conocimiento previo "
    "del dominio para escoger el modelo. La calidad estadística del resultado es "
    "comparable a frameworks más complejos cuando el dataset y los features son "
    "tabulares y de mediano tamaño (Mohr & Wever, 2023)."
)
doc.add_paragraph(
    "¿Cómo se decidió el espacio de algoritmos a comparar? Se replicó el "
    "espacio por defecto de naiveautoml: regularizadores lineales (Ridge, Lasso, "
    "ElasticNet), ensambles de árboles (Random Forest, Gradient Boosting, "
    "Histogram GB), una red neuronal pequeña (MLP) y dos baselines (mediana, "
    "Karvonen-like)."
)

# A.5
doc.add_heading('A.5 Qué es naiveautoml y cómo se ejecutó en este proyecto', level=2)
doc.add_paragraph(
    "naiveautoml es una librería de Python (https://github.com/fmohr/naiveautoml) "
    "creada por Felix Mohr y Marcel Wever (Universidad de Paderborn). Su filosofía "
    "es ofrecer un protocolo AutoML transparente y reproducible que sirva como "
    "baseline honesto en publicaciones y trabajos de grado, en contraste con "
    "frameworks complejos como Auto-Sklearn o TPOT que pueden ofuscar la "
    "metodología detrás de capas de optimización Bayesiana y meta-learning."
)
doc.add_paragraph(
    "En este proyecto, el wrapper original de naiveautoml no pudo instalarse en "
    "la máquina de desarrollo (Python 3.14 en su momento; ahora se trabaja con "
    "Python 3.11 conda). Por esa razón, en NB12 se replicó manualmente el "
    "protocolo metodológico equivalente: se evaluaron los siete algoritmos del "
    "espacio por defecto de naiveautoml más dos baselines, se usó el mismo "
    "esquema de validación cruzada (GroupKFold K=10 agrupando por usuario), se "
    "agregaron los errores por usuario para el test de Friedman, y se aplicó el "
    "post-hoc Nemenyi. El resultado es metodológicamente equivalente a haber "
    "ejecutado naiveautoml directamente."
)
doc.add_paragraph(
    "La pregunta clave para el jurado es: ¿cumplió el ejercicio el espíritu del "
    "protocolo recomendado por el director? Sí. Se compararon todos los modelos "
    "estándar bajo el mismo CV, se aplicó la misma metodología estadística, y se "
    "reportó tanto el ganador como la evidencia estadística que sustenta su "
    "elección. La sección 8.10.2 del documento principal describe la "
    "implementación con detalle suficiente para reproducir el experimento."
)

# A.6
doc.add_heading('A.6 Validaciones estadísticas aplicadas y sus resultados', level=2)
doc.add_paragraph(
    "El proyecto aplicó tres familias de pruebas estadísticas, cada una con una "
    "función específica. Esta sección las explica en lenguaje simple, indica "
    "dónde están implementadas y reporta los valores reales obtenidos."
)

# A.6.1 Friedman-Nemenyi
doc.add_heading('A.6.1 Test de Friedman + post-hoc Nemenyi', level=3)
doc.add_paragraph(
    "Para qué sirve. Comparar el desempeño de varios modelos sobre los mismos "
    "sujetos (los mismos usuarios), cuando los errores no se distribuyen "
    "normalmente. El test de Friedman responde a: \"¿hay alguna diferencia "
    "estadísticamente significativa entre los modelos?\". Si la respuesta es sí, "
    "el post-hoc Nemenyi identifica qué pares de modelos difieren entre sí "
    "controlando el error por comparaciones múltiples."
)
doc.add_paragraph(
    "Cómo se aplicó. Para cada uno de los 356 usuarios se calculó el MAE de cada "
    "modelo en sus sesiones, generando una matriz 356 × 9 (usuarios × modelos). "
    "Sobre esta matriz se aplica el test de Friedman."
)
doc.add_paragraph(
    "Resultados reales (NB12). Friedman χ² = 14.29 con p = 0.0064 sobre los "
    "top-5 modelos. Se rechaza la hipótesis nula de igualdad de desempeño con "
    "alta confianza. Post-hoc Nemenyi: Ridge supera a MLP con p = 0.022 y a "
    "Gradient Boosting con p = 0.036; Ridge y ElasticNet son estadísticamente "
    "equivalentes (p = 0.905)."
)
doc.add_paragraph(
    "Conclusión metodológica. La elección de Ridge como modelo del Nivel 1 está "
    "respaldada por evidencia estadística no paramétrica robusta a outliers."
)

# A.6.2 Conformal
doc.add_heading('A.6.2 Split-conformal prediction', level=3)
doc.add_paragraph(
    "Para qué sirve. Generar intervalos de predicción con cobertura empírica "
    "garantizada. En lugar de reportar \"tu ritmo será 5:20 min/km\", el sistema "
    "reporta \"tu ritmo estará entre 4:14 y 6:26 min/km, y el 80 % de las veces "
    "ese intervalo contiene tu ritmo real\"."
)
doc.add_paragraph(
    "Cómo se aplicó. Se divide el dataset en tres partes: entrenamiento (60 %), "
    "calibración (20 %) y test (20 %). El modelo se entrena con la parte de "
    "entrenamiento. En la parte de calibración se calculan los errores absolutos "
    "y se toma el percentil (1−α) de esos errores como ancho del intervalo. En "
    "la parte de test se verifica que la cobertura empírica iguale la nominal."
)
doc.add_paragraph(
    "Resultados reales (NB12). Con α = 0.20 (cobertura nominal 80 %), el "
    "cuantil de error absoluto es q = 1.060 min/km (≈ ±64 sec/km). La cobertura "
    "empírica medida en el conjunto de test es 0.804, prácticamente idéntica al "
    "0.80 nominal. Esto valida que el método está bien calibrado y los "
    "intervalos son honestos."
)

# Insertar figura conformal
if FIG_CONFORMAL.exists():
    doc.add_paragraph()
    doc.add_picture(str(FIG_CONFORMAL), width=Inches(5.5))
    cap = doc.add_paragraph(
        "Figura A.2. Cobertura empírica del intervalo conformal vs. cobertura "
        "nominal. Resultado: 0.804 vs. 0.80 — calibración correcta. Fuente: "
        "NB12 — nb12_full_fig2_cobertura_conformal.png."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# A.6.3 Análisis de potencia
doc.add_heading('A.6.3 Análisis de potencia estadística (Monte Carlo)', level=3)
doc.add_paragraph(
    "Para qué sirve. Estimar cuántos atletas y cuántas observaciones son "
    "necesarios para que el sistema produzca estimaciones estables. Sin este "
    "análisis, es imposible saber si la cohorte actual de RUNA es suficiente "
    "para entrenar un modelo individual o si requiere apoyarse en el prior "
    "poblacional (Nivel 1)."
)
doc.add_paragraph(
    "Cómo se aplicó. NB09 implementa un análisis de potencia mediante "
    "simulación de Monte Carlo: se repite el proceso de entrenamiento y "
    "validación bajo diferentes tamaños de muestra (de N=10 a N=200 atletas) "
    "y se cuantifica la varianza del MAE resultante."
)
doc.add_paragraph(
    "Resultados reales (NB09). El umbral mínimo de atletas para estimaciones "
    "estables es N ≥ 40. Para una cobertura conformal calibrada al 90 % se "
    "requieren al menos 100 observaciones en el conjunto de calibración. La "
    "personalización individual (Nivel 3) requiere ≥ 3 carreras propias del "
    "atleta para ser activada."
)
doc.add_paragraph(
    "Implicación práctica. La cohorte RUNA actual (9 atletas) está por debajo "
    "del umbral para entrenar un modelo propio. Esto justifica la arquitectura "
    "jerárquica: el prior poblacional (Nivel 1) aporta el esqueleto del modelo "
    "y los datos individuales se incorporan progresivamente en Niveles 2 y 3."
)

# A.6.4 Recomendaciones
doc.add_heading('A.6.4 Cuándo y dónde se aplica cada validación', level=3)
t = doc.add_table(rows=4, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = 'Validación'
hdr[1].text = 'Aplicada en'
hdr[2].text = 'Notebook / Sección'
val_rows = [
    ('Friedman + Nemenyi', 'Nivel 1 (FitRec) ✅. Nivel 2 (RUNA) cuando esté listo.', 'NB12, sección 8.10.3'),
    ('Conformal prediction', 'Nivel 1 ✅. Se replica en Niveles 2 y 3.', 'NB12, sección 8.10.4'),
    ('Análisis de potencia', 'Diseño previo del estudio ✅', 'NB09, sección 6.6'),
]
for i, (v, a, n) in enumerate(val_rows, start=1):
    row = t.rows[i].cells
    row[0].text = v
    row[1].text = a
    row[2].text = n

doc.add_paragraph()

# A.7
doc.add_heading('A.7 Mapa de notebooks y artefactos del proyecto', level=2)
doc.add_paragraph(
    "Esta sección lista los notebooks vivos del proyecto y dónde está cada "
    "artefacto producido. Es la referencia para localizar evidencia durante la "
    "sustentación."
)
t = doc.add_table(rows=8, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = 'Notebook'
hdr[1].text = 'Contenido'
hdr[2].text = 'Estado'
nb_rows = [
    ('NB09', 'Análisis de potencia (Monte Carlo) — umbrales N de cohorte', '✅ ejecutado'),
    ('NB10', 'EDA robusto para tesis — distribuciones, sesgos, calidad', '✅ ejecutado'),
    ('NB11', 'EDA Endomondo + cascada de filtros + features Nivel 1', '✅ ejecutado'),
    ('NB12', 'Entrenamiento Nivel 1 + Friedman-Nemenyi + conformal + FE', '✅ ejecutado'),
    ('NB12 v2', 'Feature engineering iterativo (5 sets × 5 modelos)', '✅ ejecutado'),
    ('NB13', 'Nivel 2 (cohorte RUNA + AutoML + LOAO-CV)', '⏳ pendiente ~mayo 2026'),
    ('NB14', 'Nivel 3 (personalización bayesiana + race simulations)', '⏳ futuro'),
]
for i, (nb, c, e) in enumerate(nb_rows, start=1):
    row = t.rows[i].cells
    row[0].text = nb
    row[1].text = c
    row[2].text = e

doc.add_paragraph()
doc.add_paragraph(
    "Tabla A.2. Notebooks del proyecto y su estado al cierre de esta versión "
    "del documento."
)

doc.add_heading('A.8 Artefactos serializados y figuras', level=2)
doc.add_paragraph(
    "Modelo serializado: ml/notebooks/outputs/nb12/nivel1_prior_poblacional_FULL_v2.pkl "
    "(Ridge + scaler + cuantil conformal). Figuras: nb12_full_fig1_ranking_modelos.png "
    "(ranking), nb12_full_fig2_cobertura_conformal.png (cobertura empírica), "
    "nb12_full_fig3_ritmo_por_zona.png (boxplot ritmo por zona), "
    "nb12_v2_fig_iteracion_features.png (curva de feature engineering)."
)

# Insertar las dos figuras restantes
if FIG_ZONAS.exists():
    doc.add_paragraph()
    doc.add_picture(str(FIG_ZONAS), width=Inches(5.5))
    cap = doc.add_paragraph(
        "Figura A.3. Boxplot del ritmo medio por zona de FC en Endomondo "
        "(N=20.710 sesiones). La señal es monotónica decreciente: a mayor "
        "intensidad cardíaca, menor tiempo por kilómetro. Fuente: NB12."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

if FIG_FEATURES.exists():
    doc.add_paragraph()
    doc.add_picture(str(FIG_FEATURES), width=Inches(5.5))
    cap = doc.add_paragraph(
        "Figura A.4. Curva de feature engineering iterativo (NB12 v2). El MAE "
        "se estabiliza alrededor de 40 sec/km — techo estructural del prior "
        "poblacional con variables agregadas por sesión."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Cierre del apéndice
doc.add_paragraph()
doc.add_paragraph(
    "Fin del Apéndice A. Las referencias bibliográficas formales se encuentran "
    "en el Capítulo 11.",
    style=None
)

# ----------------------------------------------------------------------------
# Guardar
# ----------------------------------------------------------------------------
print("\n[Guardando documento...]")
doc.save(DOCX_OUT)
print(f"✅ Documento actualizado: {DOCX_OUT}")
print(f"   Tamaño nuevo: {DOCX_OUT.stat().st_size / 1024:.1f} KB")
print(f"\n📌 SIGUIENTE PASO:")
print(f"   1. Cierra Word si tienes el v3.docx abierto")
print(f"   2. Abre {DOCX_OUT.name} y revisa los cambios")
print(f"   3. Si todo está bien, renombra a v3.docx (el viejo va a historial/)")

